from pathlib import Path
from docx import Document
import sys
sys.path.append(str(Path(__file__).resolve().parent))

from utils.docx_utils import DocxUtils
from doc_handlers.google_docs_handler import GoogleDocsHandler
from doc_handlers.document_handler_interface import DocumentHandler

class DocCompiler:
    def __init__(self):
        self.docx_utils = DocxUtils()
        self.google_handler = GoogleDocsHandler()
    
    #TODO: Add other kinds of documents in the future
    def get_handler(self, link: str) -> DocumentHandler:
        """
        Determine the appropriate handler for a given document link.

        Args:
            link (str): The URL of the document.

        Returns:
            str: The handler name or function associated with the document type.
                Returns None if no handler is found.

        Behavior:
            - Currently supports Google Docs links only.
            - Can be extended in the future to handle other document types.

        Example:
            handler = get_handler("https://docs.google.com/document/d/ABC123XYZ456/edit")
            doc_id = handler.extract_doc_id(link)
            handler.download_doc(doc_id, temp_doc_path)
        """
        handler = None
        if "docs.google.com/document/d/" in link:
            handler = self.google_handler
        return handler

    def compile_to_docx(self, origem: Path, final_path: Path, temp_dir: Path):
        """
        Compile multiple linked documents from a source DOCX into a single DOCX file.

        Args:
            origem (Path): Path to the source DOCX containing hyperlinks.
            final_path (Path): Path where the compiled DOCX will be saved.
            temp_dir (Path): Temporary directory for downloaded documents.

        Behavior:
            - Extracts hyperlinks from the source DOCX using DocxUtils.
            - Determines the handler for each link (currently supports Google Docs only).
            - Downloads each document as DOCX and copies its content (text, formatting, images) into the final document.
            - Adds page breaks between documents.
            - Handles errors gracefully, adding a message to the final document if a link cannot be processed.
            - Saves the compiled DOCX at the specified path.

        Example:
            compile_to_docx(Path("source.docx"), Path("compiled.docx"), Path("temp_docs"))
        """
        temp_dir.mkdir(exist_ok=True)

        docx_utils = DocxUtils()

        links = docx_utils.extrac_links_from_doc_per_line(origem)
        print(f"{len(links)} links found.")
        compiled_doc = Document()
        compiled_doc.add_heading("Compiled document", level=1)

        for idx, link in enumerate(links, start=1):
            handler = self.get_handler(link)
            if handler is None:
                print("File type is not supported. Currently only works with Google Docs")
                continue

            print(f"\nProcessing link {idx}/{len(links)}: {link}")
            doc_id = handler.extract_doc_id(link)
            if not doc_id:
                print(f"Document ID not found: {link}")
                continue
            try:
                temp_doc_path = temp_dir / f"{doc_id}.docx"
                handler.download_doc(doc_id, temp_doc_path)
                print(f"Doc downloaded: {temp_doc_path.name}")

                docx_utils.copy_docx_with_pictures(temp_doc_path, compiled_doc)
                print(f"Doc content copied to compiled docx file.")

                if idx < len(links):
                    compiled_doc.add_page_break()

            except Exception as e:
                print(f"Error proccessing file {link}: {e}")
                compiled_doc.add_paragraph(f"Not possible to proccess file {link}: {e}")

        compiled_doc.save(final_path)
        print(f"\nFinished! Compiled document saved: {final_path}")
