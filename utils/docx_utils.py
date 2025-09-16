from docx import Document
from docx.shared import Cm
from pathlib import Path

class DocxUtils:
    def __init__(self):
        pass

    def extrac_links_from_doc_per_line(self, docx_path: Path) -> list[str]:
        """
        Extract all hyperlinks from a DOCX file in the order they appear.

        Args:
            docx_path (Path): Path to the source DOCX file.

        Returns:
            list[str]: List of URLs corresponding to the hyperlinks found in the document.
        
        Example:
            urls = self.extrac_links_from_doc_per_line(Path("source.docx"))
            print(urls)  # ['https://example.com', 'https://another.com']
        """
        doc = Document(docx_path)
        urls = []

        rel_dict = {rel.rId: rel.target_ref for rel in doc.part.rels.values() if "hyperlink" in rel.reltype}

        for line in doc.paragraphs:
            for hyperlink in line._p.findall('.//w:hyperlink', doc.part._element.nsmap):
                rId = hyperlink.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id')
                if rId and rId in rel_dict:
                    urls.append(rel_dict[rId])

        return urls

    def copy_docx_with_pictures(self, src_path: Path, dest_doc: Document):
        """
        Copy the content of a DOCX file to another DOCX document, preserving text formatting and inline images.

        Args:
            src_path (Path): Path to the source DOCX file.
            dest_doc (Document): A python-docx Document object where the content will be copied.

        Behavior:
            - Copies all paragraphs from the source document to the destination.
            - Preserves text formatting (bold, italic, underline, font size, font name).
            - Copies inline images from the source document and inserts them into the destination.
            - Temporarily saves images to insert them, then deletes the temporary files.

        Example:
            dest_doc = Document()
            copy_docx_with_pictures(Path("source.docx"), dest_doc)
            dest_doc.save("compiled.docx")
        """
        src_doc = Document(src_path)

        for line in src_doc.paragraphs:
            # Novo parágrafo no destino
            new_para = dest_doc.add_paragraph(style=line.style)

            # Copia runs
            for run in line.runs:
                new_run = new_para.add_run(run.text)
                new_run.bold = run.bold
                new_run.italic = run.italic
                new_run.underline = run.underline
                new_run.font.size = run.font.size
                new_run.font.name = run.font.name

            # Copiar imagens (inline)
            for drawing in line._p.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/picture}pic'):
                blip = drawing.find('.//{http://schemas.openxmlformats.org/drawingml/2006/main}blip')
                if blip is not None:
                    rEmbed = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                    if rEmbed in src_doc.part.related_parts:
                        image_part = src_doc.part.related_parts[rEmbed]
                        image_bytes = image_part.blob

                        # Salva temporariamente
                        temp_img_path = Path("temp_image.png")
                        with open(temp_img_path, "wb") as f:
                            f.write(image_bytes)

                        # Insere imagem no documento destino
                        dest_doc.add_picture(str(temp_img_path), width=Cm(12))
                        temp_img_path.unlink()
